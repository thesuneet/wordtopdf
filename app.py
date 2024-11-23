from flask import Flask, request, render_template, send_from_directory
from werkzeug.utils import secure_filename
import os
from docx import Document
from reportlab.pdfgen import canvas

# Initialize Flask app
app = Flask(__name__)

# Configurations
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file part", 400
    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400

    # Save file
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    # Extract Metadata
    doc = Document(filepath)
    metadata = {
        "Author": doc.core_properties.author or "Unknown",
        "Title": doc.core_properties.title or "Untitled",
        "Created": doc.core_properties.created or "Not Available",
    }

    # Convert to PDF
    pdf_filename = filename.rsplit('.', 1)[0] + '.pdf'
    pdf_filepath = os.path.join(app.config['CONVERTED_FOLDER'], pdf_filename)
    convert_to_pdf(filepath, pdf_filepath)

    # Render the response with metadata and download link
    return render_template('result.html', metadata=metadata, pdf_filename=pdf_filename)

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['CONVERTED_FOLDER'], filename)

# Utility: Convert DOCX to PDF
def convert_to_pdf(docx_path, pdf_path):
    c = canvas.Canvas(pdf_path)
    doc = Document(docx_path)
    y_position = 750
    for paragraph in doc.paragraphs:
        if y_position < 50:  # Avoid writing off the page
            c.showPage()
            y_position = 750
        c.drawString(50, y_position, paragraph.text)
        y_position -= 20  # Move down for the next line
    c.save()

if __name__ == '__main__':
    app.run(debug=True)
